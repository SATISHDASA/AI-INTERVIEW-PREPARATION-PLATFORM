import io
import re
import logging

logger = logging.getLogger(__name__)

# ─── Text Extractors ──────────────────────────────────────────────────────────

def extract_text_from_pdf(file_bytes: bytes) -> str:
    try:
        import PyPDF2
        reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
        pages  = []
        for page in reader.pages:
            t = page.extract_text()
            if t:
                pages.append(t)
        return "\n".join(pages).strip()
    except ImportError:
        logger.error("PyPDF2 not installed. Run: pip install PyPDF2")
        return ""
    except Exception as e:
        logger.error(f"PDF extraction error: {e}")
        return ""

def extract_text_from_docx(file_bytes: bytes) -> str:
    try:
        from docx import Document          # python-docx package
        doc   = Document(io.BytesIO(file_bytes))
        lines = []
        # Paragraphs
        for para in doc.paragraphs:
            t = para.text.strip()
            if t:
                lines.append(t)
        # Tables (many resumes use table layouts)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    t = cell.text.strip()
                    if t and t not in lines:
                        lines.append(t)
        return "\n".join(lines)
    except ImportError:
        logger.error("python-docx not installed. Run: pip install python-docx")
        return ""
    except Exception as e:
        logger.error(f"DOCX extraction error: {e}")
        return ""

def extract_text_from_txt(file_bytes: bytes) -> str:
    for enc in ("utf-8", "latin-1", "cp1252"):
        try:
            return file_bytes.decode(enc)
        except UnicodeDecodeError:
            continue
    return file_bytes.decode("utf-8", errors="ignore")

# ─── Main Parser ──────────────────────────────────────────────────────────────

def parse_resume(file_bytes: bytes, filename: str) -> dict:
    fname = filename.lower().strip()

    if fname.endswith(".pdf"):
        text = extract_text_from_pdf(file_bytes)
    elif fname.endswith((".docx", ".doc")):
        text = extract_text_from_docx(file_bytes)
    elif fname.endswith(".txt"):
        text = extract_text_from_txt(file_bytes)
    else:
        return {"error": f"Unsupported file '{filename}'. Upload PDF, DOCX, or TXT."}

    if not text or not text.strip():
        return {
            "error": (
                "Could not extract text from the resume. "
                "If it is a scanned / image-based PDF, text extraction is not possible. "
                "Please use a text-based PDF or DOCX file."
            )
        }

    skills     = extract_skills(text)
    experience = estimate_experience(text)
    education  = extract_education(text)
    name       = extract_name(text)

    return {
        "raw_text":         text[:5000],
        "skills":           skills,
        "experience_years": experience,
        "education":        education,
        "name":             name,
        "char_count":       len(text),
        "summary": (
            f"Experience: {experience} yr(s) | "
            f"Education: {education[0] if education else 'N/A'} | "
            f"Skills: {', '.join(skills[:8])}"
        ),
    }

# ─── Helpers ──────────────────────────────────────────────────────────────────

TECH_SKILLS = [
    # Languages
    "Python","Java","JavaScript","TypeScript","C++","C#","C","Go","Golang",
    "Rust","Kotlin","Swift","PHP","Ruby","Scala","R","MATLAB","Perl","Bash","PowerShell",
    # Frontend
    "React","React.js","Angular","Vue","Vue.js","Next.js","Nuxt.js",
    "HTML","CSS","Tailwind","Bootstrap","jQuery","Redux","Webpack","Vite",
    # Backend
    "Node.js","Express","Django","Flask","FastAPI","Spring Boot","Spring",
    "Laravel","Rails","ASP.NET","GraphQL","REST API","Microservices",
    # Data / AI / ML
    "Machine Learning","Deep Learning","NLP","Computer Vision",
    "PyTorch","TensorFlow","Keras","Scikit-learn","Pandas","NumPy",
    "Matplotlib","Seaborn","OpenCV","Spark","Hadoop","MLOps",
    "LangChain","Hugging Face","BERT","GPT","Data Science","Data Analysis",
    # Databases
    "SQL","MySQL","PostgreSQL","MongoDB","SQLite","Redis","Elasticsearch",
    "Cassandra","DynamoDB","Oracle","Firebase","Supabase",
    # Cloud / DevOps
    "AWS","GCP","Azure","Google Cloud","Docker","Kubernetes","Terraform",
    "Ansible","Jenkins","CI/CD","GitHub Actions","Linux","Unix","Heroku","Vercel",
    # Tools
    "Git","GitHub","GitLab","Jira","Agile","Scrum","Selenium",
    "Pytest","Jest","Postman","Figma","Tableau","Power BI",
]

def extract_skills(text: str) -> list:
    found = []
    seen  = set()
    for skill in TECH_SKILLS:
        if re.search(rf"\b{re.escape(skill)}\b", text, re.IGNORECASE):
            key = skill.lower().replace(".", "").replace(" ", "")
            if key not in seen:
                seen.add(key)
                found.append(skill)
    return found[:25]

def estimate_experience(text: str) -> str:
    patterns = [
        r"(\d+)\+?\s*years?\s+of\s+(?:professional\s+|work\s+)?experience",
        r"(\d+)\+?\s*years?\s+experience",
        r"experience[:\s]+(\d+)\+?\s*years?",
        r"(\d+)\+?\s*yrs?\s+(?:of\s+)?experience",
    ]
    for p in patterns:
        m = re.search(p, text, re.IGNORECASE)
        if m:
            return m.group(1)
    # Derive from year ranges in text
    years = sorted(set(int(y) for y in re.findall(r"\b(20\d{2}|19\d{2})\b", text)))
    if len(years) >= 2:
        span = years[-1] - years[0]
        if 1 <= span <= 40:
            return str(span)
    return "Fresher"

def extract_education(text: str) -> list:
    patterns = [
        r"B\.?Tech", r"B\.?E\.?", r"B\.?Sc", r"B\.?S\.?",
        r"M\.?Tech", r"M\.?E\.?", r"M\.?Sc", r"M\.?S\.?",
        r"MCA", r"BCA", r"Bachelor(?:'s)?", r"Master(?:'s)?",
        r"PhD", r"Ph\.D", r"Doctorate", r"MBA", r"B\.?Com", r"M\.?Com",
        r"Diploma", r"Associate",
    ]
    found = []
    seen  = set()
    for pat in patterns:
        m = re.search(pat, text, re.IGNORECASE)
        if m:
            key = m.group(0).lower().replace(".", "")
            if key not in seen:
                seen.add(key)
                found.append(m.group(0))
    return found if found else ["Not specified"]

def extract_name(text: str) -> str:
    for line in text.split("\n")[:6]:
        line = line.strip()
        if re.match(r"^[A-Za-z]+([\s][A-Za-z]+){1,3}$", line) and len(line) < 50:
            return line
    return "Candidate"
